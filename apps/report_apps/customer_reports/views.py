import datetime
from datetime import datetime, timedelta
from io import BytesIO

import matplotlib.pyplot as plt
from dateutil.relativedelta import relativedelta
from django.http import HttpResponse
from docx import Document
from docx.shared import Inches
from drf_yasg import openapi
from drf_yasg.utils import swagger_auto_schema
from openpyxl import Workbook
from openpyxl.styles import Alignment
from openpyxl.utils import get_column_letter
from rest_framework import status
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework.viewsets import ViewSet

from apps.qsystem.models import Customer
from apps.booking.models import Booking


class BookingReportViewSet(ViewSet):

    @action(detail=False, methods=['get'])
    @swagger_auto_schema(manual_parameters=[
        openapi.Parameter(
            name='year',
            in_=openapi.IN_QUERY,
            type=openapi.TYPE_STRING,
            format=openapi.FORMAT_INT64,
            required=True,
            description='Date in the format "YYYY"')
    ],
    operation_summary="Отчет предварительных бронирований за год. WORD",
    operation_description="Отчет в формате Word, количество предварительных бронирований за год"
    )

    def get_monthly_booking_statistics(self, request):
        year = request.query_params.get('year')

        try:
            year = int(year)
        except ValueError:
            return Response({'error': 'Invalid year format'}, status=status.HTTP_400_BAD_REQUEST)

        start_date = datetime(year, 1, 1)
        end_date = start_date + relativedelta(years=1) - timedelta(days=1)

        bookings = Booking.objects.filter(date__range=(start_date, end_date))

        monthly_bookings = [0] * 12

        for month in range(1, 13):
            month_start = datetime(year, month, 1)
            month_end = month_start + relativedelta(months=1) - timedelta(days=1)
            month_bookings = bookings.filter(date__range=(month_start, month_end))
            monthly_bookings[month-1] = month_bookings.count()

        x = ['January', 'February', 'March', 'April', 'May', 'June', 'July', 'August', 'September', 'October', 'November', 'December']
        y = monthly_bookings

        plt.figure(figsize=(8, 6))
        plt.plot(x, y)

        plt.xlabel('Month')
        plt.ylabel('Number of Bookings')
        plt.title(f'Monthly Booking Statistics - {year}')

        chart_file = 'chart.png'
        plt.savefig(chart_file)

        document = Document()
        document.add_paragraph(f'Year: {year}')
        document.add_picture(chart_file, width=Inches(6), height=Inches(4))

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename=booking_statistics_{year}.docx'
        document.save(response)

        return response



class CustomersCancelledExcelViewSet(ViewSet):
    @action(detail=False, methods=['get'])
    @swagger_auto_schema(manual_parameters=[
        openapi.Parameter(
            name='year',
            in_=openapi.IN_QUERY,
            type=openapi.TYPE_STRING,
            format=openapi.FORMAT_INT64,
            required=True,
            description='Date in the format "YYYY"')
    ],
    operation_summary="Отчет отмененных талонов за год. EXCEL",
    operation_description="Отчет в формате Excel, отмененные талоны за год"
    )
    def get_year_report(self, request):

        year = request.query_params.get('year')


        customers = self.get_customers_by_year(year)

      
        wb = Workbook()
        ws = wb.active
        ws.title = 'Hourly Report'

       
        ws['A1'] = 'ID'
        ws['B1'] = 'Номер талона'
        ws['C1'] = 'Номер паспорта'
        ws['D1'] = 'Категория'
        ws['E1'] = 'Дата'
        ws['F1'] = 'Филиал'


      
        row = 2
        for customer in customers:

            ws.cell(row=row, column=1, value=customer.id)
            ws.cell(row=row, column=2, value=customer.ticket_number)
            ws.cell(row=row, column=3, value=customer.pasport)
            ws.cell(row=row, column=4, value=customer.category)
            ws.cell(row=row, column=5, value=customer.created_at.strftime('%Y-%m-%d'))
            ws.cell(row=row, column=6, value=customer.queue.branch.name)

            for column in range(1, 5):
                column_letter = get_column_letter(column)
                cell = ws[column_letter + str(row)]
                ws.column_dimensions[column_letter].width = max(ws.column_dimensions[column_letter].width, len(str(cell.value)))

            row += 1

      
        for row in ws.iter_rows(min_row=1, max_row=ws.max_row):
            for cell in row:
                cell.alignment = Alignment(horizontal='center')

        
        output = BytesIO()
        wb.save(output)
        output.seek(0)

        
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename=customers_day_report__{year}.xlsx'
        response.write(output.getvalue())

        return response

    
    @action(detail=False, methods=['get'])
    @swagger_auto_schema(manual_parameters=[
        openapi.Parameter(
            name='date',
            in_=openapi.IN_QUERY,
            type=openapi.TYPE_STRING,
            format=openapi.FORMAT_DATE,
            required=True,
            description='Date in the format "YYYY-MM-DD"'
        )
    ],
    operation_summary="Отчет отмененных талонов за день. EXCEL",
    operation_description="Отчет в формате Excel, отмененные талоны за день"
    )
    def get_day_report(self, request):

        date = request.query_params.get('date')


        customers = self.get_customers_by_day(date)


        wb = Workbook()
        ws = wb.active
        ws.title = 'Hourly Report'


        ws['A1'] = 'ID'
        ws['B1'] = 'Номер талона'
        ws['C1'] = 'Время обслуживания(с)'
        ws['D1'] = 'Начало обслуживания'
        ws['E1'] = 'Конец обслуживания'
        ws['F1'] = 'Время ожидания'
        ws['G1'] = 'Номер окна'
        ws['H1'] = 'Номер паспорта'
        ws['I1'] = 'Категория'
        ws['J1'] = 'Оператор'
        ws['K1'] = 'Дата'
        ws['L1'] = 'Филиал'


        row = 2
        for customer in customers:

            served_start = customer.served_start.replace(tzinfo=None)
            served_at = customer.served_at.replace(tzinfo=None)

            ws.cell(row=row, column=1, value=customer.id)
            ws.cell(row=row, column=2, value=customer.ticket_number)
            ws.cell(row=row, column=3, value=(served_at - served_start).total_seconds())
            ws.cell(row=row, column=4, value=served_start.strftime('%H:%M:%S'))  
            ws.cell(row=row, column=5, value=served_at.strftime('%H:%M:%S'))
            ws.cell(row=row, column=6, value=(customer.served_start - customer.created_at).total_seconds())
            ws.cell(row=row, column=7, value=customer.window.number)
            ws.cell(row=row, column=8, value=customer.pasport)
            ws.cell(row=row, column=9, value=customer.category)
            ws.cell(row=row, column=10, value=customer.operator.username)
            ws.cell(row=row, column=11, value=customer.created_at.strftime('%Y-%m-%d'))
            ws.cell(row=row, column=12, value=customer.queue.branch.name)

            for column in range(1, 11):
                column_letter = get_column_letter(column)
                cell = ws[column_letter + str(row)]
                ws.column_dimensions[column_letter].width = max(ws.column_dimensions[column_letter].width, len(str(cell.value)))

            row += 1


        for row in ws.iter_rows(min_row=1, max_row=ws.max_row):
            for cell in row:
                cell.alignment = Alignment(horizontal='center')


        output = BytesIO()
        wb.save(output)
        output.seek(0)


        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename=customers_day_report_{date}.xlsx'
        response.write(output.getvalue())

        return response
    

    @action(detail=False, methods=['get'])
    @swagger_auto_schema(manual_parameters=[
        openapi.Parameter(
            name='start_date',
            in_=openapi.IN_QUERY,
            type=openapi.TYPE_STRING,
            format=openapi.FORMAT_DATE,
            required=True,
            description='Date in the format "YYYY-MM-DD"'
        ),
        openapi.Parameter(
            name='end_date',
            in_=openapi.IN_QUERY,
            type=openapi.TYPE_STRING,
            format=openapi.FORMAT_DATE,
            required=True,
            description='Date in the format "YYYY-MM-DD"'
        )
    ],
    operation_summary="Отчет отмененных талонов за промежуток дней. EXCEL",
    operation_description="Отчет в формате Excel, отмененные талоны за выбранный промежуток дней"
    )
    def get_days_report(self, request):

        start_date = request.query_params.get('start_date')
        end_date = request.query_params.get('end_date')


        customers = self.get_customers_by_days(start_date, end_date)


        wb = Workbook()
        ws = wb.active
        ws.title = 'Hourly Report'


        ws['A1'] = 'ID'
        ws['B1'] = 'Номер талона'
        ws['C1'] = 'Время обслуживания(с)'
        ws['D1'] = 'Начало обслуживания'
        ws['E1'] = 'Конец обслуживания'
        ws['F1'] = 'Время ожидания'
        ws['G1'] = 'Номер окна'
        ws['H1'] = 'Номер паспорта'
        ws['I1'] = 'Категория'
        ws['J1'] = 'Оператор'
        ws['K1'] = 'Дата'
        ws['L1'] = 'Филиал'


        row = 2
        for customer in customers:

            served_start = customer.served_start.replace(tzinfo=None)
            served_at = customer.served_at.replace(tzinfo=None)

            ws.cell(row=row, column=1, value=customer.id)
            ws.cell(row=row, column=2, value=customer.ticket_number)
            ws.cell(row=row, column=3, value=(served_at - served_start).total_seconds())
            ws.cell(row=row, column=4, value=served_start.strftime('%H:%M:%S')) 
            ws.cell(row=row, column=5, value=served_at.strftime('%H:%M:%S'))
            ws.cell(row=row, column=6, value=(customer.served_start - customer.created_at).total_seconds())
            ws.cell(row=row, column=7, value=customer.window.number)
            ws.cell(row=row, column=8, value=customer.pasport)
            ws.cell(row=row, column=9, value=customer.category)
            ws.cell(row=row, column=10, value=customer.operator.username)
            ws.cell(row=row, column=11, value=customer.created_at.strftime('%Y-%m-%d'))
            ws.cell(row=row, column=12, value=customer.queue.branch.name)

            for column in range(1, 11):
                column_letter = get_column_letter(column)
                cell = ws[column_letter + str(row)]
                ws.column_dimensions[column_letter].width = max(ws.column_dimensions[column_letter].width, len(str(cell.value)))

            row += 1


        for row in ws.iter_rows(min_row=1, max_row=ws.max_row):
            for cell in row:
                cell.alignment = Alignment(horizontal='center')


        output = BytesIO()
        wb.save(output)
        output.seek(0)

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename=customers_day_report_{start_date}_{end_date}.xlsx'
        response.write(output.getvalue())

        return response
    

    def get_customers_by_days(self, start_date, end_date):
        customers = Customer.objects.filter(created_at__date__range=(start_date, end_date), is_served=False)
        return customers

    def get_customers_by_day(self, date):
        customers = Customer.objects.filter(created_at__date=date, is_served=False)
        return customers
    
    def get_customers_by_year(self, year):
        customers = Customer.objects.filter(created_at__year=year, is_served=False)
        return customers


class CustomersServedExcelViewSet(ViewSet):
    @action(detail=False, methods=['get'])
    @swagger_auto_schema(manual_parameters=[
        openapi.Parameter(
            name='year',
            in_=openapi.IN_QUERY,
            type=openapi.TYPE_STRING,
            format=openapi.FORMAT_INT64,
            required=True,
            description='Date in the format "YYYY"')
    ],
    operation_summary="Отчет обслуженных талонов за год. EXCEL",
    operation_description="Отчет в формате Excel, обслуженные талоны за год"

    )
    def get_year_report(self, request):

        year = request.query_params.get('year')


        customers = self.get_customers_by_year(year)

      
        wb = Workbook()
        ws = wb.active
        ws.title = 'Hourly Report'

       
        ws['A1'] = 'ID'
        ws['B1'] = 'Номер талона'
        ws['C1'] = 'Номер паспорта'
        ws['D1'] = 'Категория'
        ws['E1'] = 'Дата'
        ws['F1'] = 'Филиал'


      
        row = 2
        for customer in customers:

            ws.cell(row=row, column=1, value=customer.id)
            ws.cell(row=row, column=2, value=customer.ticket_number)
            ws.cell(row=row, column=3, value=customer.pasport)
            ws.cell(row=row, column=4, value=customer.category)
            ws.cell(row=row, column=5, value=customer.created_at.strftime('%Y-%m-%d'))
            ws.cell(row=row, column=6, value=customer.queue.branch.name)

            for column in range(1, 5):
                column_letter = get_column_letter(column)
                cell = ws[column_letter + str(row)]
                ws.column_dimensions[column_letter].width = max(ws.column_dimensions[column_letter].width, len(str(cell.value)))

            row += 1

      
        for row in ws.iter_rows(min_row=1, max_row=ws.max_row):
            for cell in row:
                cell.alignment = Alignment(horizontal='center')

        
        output = BytesIO()
        wb.save(output)
        output.seek(0)

        
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename=customers_day_report__{year}.xlsx'
        response.write(output.getvalue())

        return response

    
    @action(detail=False, methods=['get'])
    @swagger_auto_schema(manual_parameters=[
        openapi.Parameter(
            name='date',
            in_=openapi.IN_QUERY,
            type=openapi.TYPE_STRING,
            format=openapi.FORMAT_DATE,
            required=True,
            description='Date in the format "YYYY-MM-DD"'
        )
    ],
    operation_summary="Отчет обслуженных талонов за день. EXCEL",
    operation_description="Отчет в формате Excel, обслуженные талоны за выбранный день"
    )
    def get_day_report(self, request):

        date = request.query_params.get('date')


        customers = self.get_customers_by_day(date)


        wb = Workbook()
        ws = wb.active
        ws.title = 'Hourly Report'


        ws['A1'] = 'ID'
        ws['B1'] = 'Номер талона'
        ws['C1'] = 'Время обслуживания(с)'
        ws['D1'] = 'Начало обслуживания'
        ws['E1'] = 'Конец обслуживания'
        ws['F1'] = 'Время ожидания'
        ws['G1'] = 'Номер окна'
        ws['H1'] = 'Номер паспорта'
        ws['I1'] = 'Категория'
        ws['J1'] = 'Оператор'
        ws['K1'] = 'Дата'
        ws['L1'] = 'Филиал'


        row = 2
        for customer in customers:

            served_start = customer.served_start.replace(tzinfo=None)
            served_at = customer.served_at.replace(tzinfo=None)

            ws.cell(row=row, column=1, value=customer.id)
            ws.cell(row=row, column=2, value=customer.ticket_number)
            ws.cell(row=row, column=3, value=(served_at - served_start).total_seconds())
            ws.cell(row=row, column=4, value=served_start.strftime('%H:%M:%S'))  
            ws.cell(row=row, column=5, value=served_at.strftime('%H:%M:%S'))
            ws.cell(row=row, column=6, value=(customer.served_start - customer.created_at).total_seconds())
            ws.cell(row=row, column=7, value=customer.window.number)
            ws.cell(row=row, column=8, value=customer.pasport)
            ws.cell(row=row, column=9, value=customer.category)
            ws.cell(row=row, column=10, value=customer.operator.username)
            ws.cell(row=row, column=11, value=customer.created_at.strftime('%Y-%m-%d'))
            ws.cell(row=row, column=12, value=customer.queue.branch.name)

            for column in range(1, 11):
                column_letter = get_column_letter(column)
                cell = ws[column_letter + str(row)]
                ws.column_dimensions[column_letter].width = max(ws.column_dimensions[column_letter].width, len(str(cell.value)))

            row += 1


        for row in ws.iter_rows(min_row=1, max_row=ws.max_row):
            for cell in row:
                cell.alignment = Alignment(horizontal='center')


        output = BytesIO()
        wb.save(output)
        output.seek(0)


        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename=customers_day_report_{date}.xlsx'
        response.write(output.getvalue())

        return response
    

    @action(detail=False, methods=['get'])
    @swagger_auto_schema(manual_parameters=[
        openapi.Parameter(
            name='start_date',
            in_=openapi.IN_QUERY,
            type=openapi.TYPE_STRING,
            format=openapi.FORMAT_DATE,
            required=True,
            description='Date in the format "YYYY-MM-DD"'
        ),
        openapi.Parameter(
            name='end_date',
            in_=openapi.IN_QUERY,
            type=openapi.TYPE_STRING,
            format=openapi.FORMAT_DATE,
            required=True,
            description='Date in the format "YYYY-MM-DD"'
        )
    ],
    operation_summary="Отчет обслуженных талонов за промежуток дней. EXCEL",
    operation_description="Отчет в формате Excel, обслуженные талоны за выбранный промежуток дней"
    )
    def get_days_report(self, request):

        start_date = request.query_params.get('start_date')
        end_date = request.query_params.get('end_date')


        customers = self.get_customers_by_days(start_date, end_date)


        wb = Workbook()
        ws = wb.active
        ws.title = 'Hourly Report'


        ws['A1'] = 'ID'
        ws['B1'] = 'Номер талона'
        ws['C1'] = 'Время обслуживания(с)'
        ws['D1'] = 'Начало обслуживания'
        ws['E1'] = 'Конец обслуживания'
        ws['F1'] = 'Время ожидания'
        ws['G1'] = 'Номер окна'
        ws['H1'] = 'Номер паспорта'
        ws['I1'] = 'Категория'
        ws['J1'] = 'Оператор'
        ws['K1'] = 'Дата'
        ws['L1'] = 'Филиал'


        row = 2
        for customer in customers:

            served_start = customer.served_start.replace(tzinfo=None)
            served_at = customer.served_at.replace(tzinfo=None)

            ws.cell(row=row, column=1, value=customer.id)
            ws.cell(row=row, column=2, value=customer.ticket_number)
            ws.cell(row=row, column=3, value=(served_at - served_start).total_seconds())
            ws.cell(row=row, column=4, value=served_start.strftime('%H:%M:%S')) 
            ws.cell(row=row, column=5, value=served_at.strftime('%H:%M:%S'))
            ws.cell(row=row, column=6, value=(customer.served_start - customer.created_at).total_seconds())
            ws.cell(row=row, column=7, value=customer.window.number)
            ws.cell(row=row, column=8, value=customer.pasport)
            ws.cell(row=row, column=9, value=customer.category)
            ws.cell(row=row, column=10, value=customer.operator.username)
            ws.cell(row=row, column=11, value=customer.created_at.strftime('%Y-%m-%d'))
            ws.cell(row=row, column=12, value=customer.queue.branch.name)

            for column in range(1, 11):
                column_letter = get_column_letter(column)
                cell = ws[column_letter + str(row)]
                ws.column_dimensions[column_letter].width = max(ws.column_dimensions[column_letter].width, len(str(cell.value)))

            row += 1


        for row in ws.iter_rows(min_row=1, max_row=ws.max_row):
            for cell in row:
                cell.alignment = Alignment(horizontal='center')


        output = BytesIO()
        wb.save(output)
        output.seek(0)

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename=customers_day_report_{start_date}_{end_date}.xlsx'
        response.write(output.getvalue())

        return response
    

    def get_customers_by_days(self, start_date, end_date):
        customers = Customer.objects.filter(created_at__date__range=(start_date, end_date), is_served=True)
        return customers

    def get_customers_by_day(self, date):
        customers = Customer.objects.filter(created_at__date=date, is_served=True)
        return customers
    
    def get_customers_by_year(self, year):
        customers = Customer.objects.filter(created_at__year=year, is_served=True)
        return customers
