from pathlib import Path

from django.http import FileResponse
from rest_framework.decorators import api_view


import datetime
import tempfile
from datetime import datetime, timedelta
from io import BytesIO

import docx
import matplotlib.pyplot as plt
from dateutil.relativedelta import relativedelta
from django.contrib.auth.models import User
from django.db.models import Count
from django.db.models.functions import ExtractMonth
from django.http import HttpResponse
from docx import Document
from docx.shared import Inches
from drf_yasg import openapi
from drf_yasg.utils import swagger_auto_schema
from openpyxl import Workbook
from openpyxl.styles import Alignment
from openpyxl.utils import get_column_letter
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from rest_framework.decorators import action
from rest_framework.views import APIView
from rest_framework.viewsets import ViewSet

from apps.booking.models import Booking
from apps.qsystem.models import Customer

from .models import OperatorDailyReport

BASE_DIR = Path(__file__).resolve().parent.parent.parent

class BookingReportViewSet(ViewSet):

    @action(detail=False, methods=['get'])
    @swagger_auto_schema(manual_parameters=[
        openapi.Parameter(
            name='year',
            in_=openapi.IN_QUERY,
            type=openapi.TYPE_STRING,
            format=openapi.FORMAT_INT64,
            required=True,
            description='Date in the format "YYYY"')
    ])

    def get_monthly_booking_statistics(self, request):
        year = request.query_params.get('year')

        try:
            year = int(year)
        except ValueError:
            return Response({'error': 'Invalid year format'}, status=status.HTTP_400_BAD_REQUEST)

        start_date = datetime(year, 1, 1)
        end_date = start_date + relativedelta(years=1) - timedelta(days=1)

        bookings = Booking.objects.filter(date__range=(start_date, end_date))

        monthly_bookings = [0] * 12

        for month in range(1, 13):
            month_start = datetime(year, month, 1)
            month_end = month_start + relativedelta(months=1) - timedelta(days=1)
            month_bookings = bookings.filter(date__range=(month_start, month_end))
            monthly_bookings[month-1] = month_bookings.count()

        x = ['January', 'February', 'March', 'April', 'May', 'June', 'July', 'August', 'September', 'October', 'November', 'December']
        y = monthly_bookings

        plt.figure(figsize=(8, 6))
        plt.plot(x, y)

        plt.xlabel('Month')
        plt.ylabel('Number of Bookings')
        plt.title(f'Monthly Booking Statistics - {year}')

        chart_file = 'chart.png'
        plt.savefig(chart_file)

        document = Document()
        document.add_paragraph(f'Year: {year}')
        document.add_picture(chart_file, width=Inches(6), height=Inches(4))

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename=booking_statistics_{year}.docx'
        document.save(response)

        return response





class CustomersCancelledExcelViewSet(ViewSet):
    @action(detail=False, methods=['get'])
    @swagger_auto_schema(manual_parameters=[
        openapi.Parameter(
            name='year',
            in_=openapi.IN_QUERY,
            type=openapi.TYPE_STRING,
            format=openapi.FORMAT_INT64,
            required=True,
            description='Date in the format "YYYY"')
    ])
    def get_year_report(self, request):

        year = request.query_params.get('year')


        customers = self.get_customers_by_year(year)

      
        wb = Workbook()
        ws = wb.active
        ws.title = 'Hourly Report'

       
        ws['A1'] = 'ID'
        ws['B1'] = 'Номер талона'
        ws['C1'] = 'Номер паспорта'
        ws['D1'] = 'Категория'
        ws['E1'] = 'Дата'
        ws['F1'] = 'Филиал'


      
        row = 2
        for customer in customers:

            ws.cell(row=row, column=1, value=customer.id)
            ws.cell(row=row, column=2, value=customer.ticket_number)
            ws.cell(row=row, column=3, value=customer.pasport)
            ws.cell(row=row, column=4, value=customer.category)
            ws.cell(row=row, column=5, value=customer.created_at.strftime('%Y-%m-%d'))
            ws.cell(row=row, column=6, value=customer.queue.branch.name)

            for column in range(1, 5):
                column_letter = get_column_letter(column)
                cell = ws[column_letter + str(row)]
                ws.column_dimensions[column_letter].width = max(ws.column_dimensions[column_letter].width, len(str(cell.value)))

            row += 1

      
        for row in ws.iter_rows(min_row=1, max_row=ws.max_row):
            for cell in row:
                cell.alignment = Alignment(horizontal='center')

        
        output = BytesIO()
        wb.save(output)
        output.seek(0)

        
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename=customers_day_report__{year}.xlsx'
        response.write(output.getvalue())

        return response

    
    @action(detail=False, methods=['get'])
    @swagger_auto_schema(manual_parameters=[
        openapi.Parameter(
            name='date',
            in_=openapi.IN_QUERY,
            type=openapi.TYPE_STRING,
            format=openapi.FORMAT_DATE,
            required=True,
            description='Date in the format "YYYY-MM-DD"'
        )
    ])
    def get_day_report(self, request):

        date = request.query_params.get('date')


        customers = self.get_customers_by_day(date)


        wb = Workbook()
        ws = wb.active
        ws.title = 'Hourly Report'


        ws['A1'] = 'ID'
        ws['B1'] = 'Номер талона'
        ws['C1'] = 'Время обслуживания(с)'
        ws['D1'] = 'Начало обслуживания'
        ws['E1'] = 'Конец обслуживания'
        ws['F1'] = 'Время ожидания'
        ws['G1'] = 'Номер окна'
        ws['H1'] = 'Номер паспорта'
        ws['I1'] = 'Категория'
        ws['J1'] = 'Оператор'
        ws['K1'] = 'Дата'
        ws['L1'] = 'Филиал'


        row = 2
        for customer in customers:

            served_start = customer.served_start.replace(tzinfo=None)
            served_at = customer.served_at.replace(tzinfo=None)

            ws.cell(row=row, column=1, value=customer.id)
            ws.cell(row=row, column=2, value=customer.ticket_number)
            ws.cell(row=row, column=3, value=(served_at - served_start).total_seconds())
            ws.cell(row=row, column=4, value=served_start.strftime('%H:%M:%S'))  
            ws.cell(row=row, column=5, value=served_at.strftime('%H:%M:%S'))
            ws.cell(row=row, column=6, value=(customer.served_start - customer.created_at).total_seconds())
            ws.cell(row=row, column=7, value=customer.window.number)
            ws.cell(row=row, column=8, value=customer.pasport)
            ws.cell(row=row, column=9, value=customer.category)
            ws.cell(row=row, column=10, value=customer.operator.username)
            ws.cell(row=row, column=11, value=customer.created_at.strftime('%Y-%m-%d'))
            ws.cell(row=row, column=12, value=customer.queue.branch.name)

            for column in range(1, 11):
                column_letter = get_column_letter(column)
                cell = ws[column_letter + str(row)]
                ws.column_dimensions[column_letter].width = max(ws.column_dimensions[column_letter].width, len(str(cell.value)))

            row += 1


        for row in ws.iter_rows(min_row=1, max_row=ws.max_row):
            for cell in row:
                cell.alignment = Alignment(horizontal='center')


        output = BytesIO()
        wb.save(output)
        output.seek(0)


        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename=customers_day_report_{date}.xlsx'
        response.write(output.getvalue())

        return response
    

    @action(detail=False, methods=['get'])
    @swagger_auto_schema(manual_parameters=[
        openapi.Parameter(
            name='start_date',
            in_=openapi.IN_QUERY,
            type=openapi.TYPE_STRING,
            format=openapi.FORMAT_DATE,
            required=True,
            description='Date in the format "YYYY-MM-DD"'
        ),
        openapi.Parameter(
            name='end_date',
            in_=openapi.IN_QUERY,
            type=openapi.TYPE_STRING,
            format=openapi.FORMAT_DATE,
            required=True,
            description='Date in the format "YYYY-MM-DD"'
        )
    ])
    def get_days_report(self, request):

        start_date = request.query_params.get('start_date')
        end_date = request.query_params.get('end_date')


        customers = self.get_customers_by_days(start_date, end_date)


        wb = Workbook()
        ws = wb.active
        ws.title = 'Hourly Report'


        ws['A1'] = 'ID'
        ws['B1'] = 'Номер талона'
        ws['C1'] = 'Время обслуживания(с)'
        ws['D1'] = 'Начало обслуживания'
        ws['E1'] = 'Конец обслуживания'
        ws['F1'] = 'Время ожидания'
        ws['G1'] = 'Номер окна'
        ws['H1'] = 'Номер паспорта'
        ws['I1'] = 'Категория'
        ws['J1'] = 'Оператор'
        ws['K1'] = 'Дата'
        ws['L1'] = 'Филиал'


        row = 2
        for customer in customers:

            served_start = customer.served_start.replace(tzinfo=None)
            served_at = customer.served_at.replace(tzinfo=None)

            ws.cell(row=row, column=1, value=customer.id)
            ws.cell(row=row, column=2, value=customer.ticket_number)
            ws.cell(row=row, column=3, value=(served_at - served_start).total_seconds())
            ws.cell(row=row, column=4, value=served_start.strftime('%H:%M:%S')) 
            ws.cell(row=row, column=5, value=served_at.strftime('%H:%M:%S'))
            ws.cell(row=row, column=6, value=(customer.served_start - customer.created_at).total_seconds())
            ws.cell(row=row, column=7, value=customer.window.number)
            ws.cell(row=row, column=8, value=customer.pasport)
            ws.cell(row=row, column=9, value=customer.category)
            ws.cell(row=row, column=10, value=customer.operator.username)
            ws.cell(row=row, column=11, value=customer.created_at.strftime('%Y-%m-%d'))
            ws.cell(row=row, column=12, value=customer.queue.branch.name)

            for column in range(1, 11):
                column_letter = get_column_letter(column)
                cell = ws[column_letter + str(row)]
                ws.column_dimensions[column_letter].width = max(ws.column_dimensions[column_letter].width, len(str(cell.value)))

            row += 1


        for row in ws.iter_rows(min_row=1, max_row=ws.max_row):
            for cell in row:
                cell.alignment = Alignment(horizontal='center')


        output = BytesIO()
        wb.save(output)
        output.seek(0)

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename=customers_day_report_{start_date}_{end_date}.xlsx'
        response.write(output.getvalue())

        return response
    

    def get_customers_by_days(self, start_date, end_date):
        customers = Customer.objects.filter(created_at__date__range=(start_date, end_date), is_served=False)
        return customers

    def get_customers_by_day(self, date):
        customers = Customer.objects.filter(created_at__date=date, is_served=False)
        return customers
    
    def get_customers_by_year(self, year):
        customers = Customer.objects.filter(created_at__year=year, is_served=False)
        return customers


class CustomersServedExcelViewSet(ViewSet):
    @action(detail=False, methods=['get'])
    @swagger_auto_schema(manual_parameters=[
        openapi.Parameter(
            name='year',
            in_=openapi.IN_QUERY,
            type=openapi.TYPE_STRING,
            format=openapi.FORMAT_INT64,
            required=True,
            description='Date in the format "YYYY"')
    ])
    def get_year_report(self, request):

        year = request.query_params.get('year')


        customers = self.get_customers_by_year(year)

      
        wb = Workbook()
        ws = wb.active
        ws.title = 'Hourly Report'

       
        ws['A1'] = 'ID'
        ws['B1'] = 'Номер талона'
        ws['C1'] = 'Номер паспорта'
        ws['D1'] = 'Категория'
        ws['E1'] = 'Дата'
        ws['F1'] = 'Филиал'


      
        row = 2
        for customer in customers:

            ws.cell(row=row, column=1, value=customer.id)
            ws.cell(row=row, column=2, value=customer.ticket_number)
            ws.cell(row=row, column=3, value=customer.pasport)
            ws.cell(row=row, column=4, value=customer.category)
            ws.cell(row=row, column=5, value=customer.created_at.strftime('%Y-%m-%d'))
            ws.cell(row=row, column=6, value=customer.queue.branch.name)

            for column in range(1, 5):
                column_letter = get_column_letter(column)
                cell = ws[column_letter + str(row)]
                ws.column_dimensions[column_letter].width = max(ws.column_dimensions[column_letter].width, len(str(cell.value)))

            row += 1

      
        for row in ws.iter_rows(min_row=1, max_row=ws.max_row):
            for cell in row:
                cell.alignment = Alignment(horizontal='center')

        
        output = BytesIO()
        wb.save(output)
        output.seek(0)

        
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename=customers_day_report__{year}.xlsx'
        response.write(output.getvalue())

        return response

    
    @action(detail=False, methods=['get'])
    @swagger_auto_schema(manual_parameters=[
        openapi.Parameter(
            name='date',
            in_=openapi.IN_QUERY,
            type=openapi.TYPE_STRING,
            format=openapi.FORMAT_DATE,
            required=True,
            description='Date in the format "YYYY-MM-DD"'
        )
    ])
    def get_day_report(self, request):

        date = request.query_params.get('date')


        customers = self.get_customers_by_day(date)


        wb = Workbook()
        ws = wb.active
        ws.title = 'Hourly Report'


        ws['A1'] = 'ID'
        ws['B1'] = 'Номер талона'
        ws['C1'] = 'Время обслуживания(с)'
        ws['D1'] = 'Начало обслуживания'
        ws['E1'] = 'Конец обслуживания'
        ws['F1'] = 'Время ожидания'
        ws['G1'] = 'Номер окна'
        ws['H1'] = 'Номер паспорта'
        ws['I1'] = 'Категория'
        ws['J1'] = 'Оператор'
        ws['K1'] = 'Дата'
        ws['L1'] = 'Филиал'


        row = 2
        for customer in customers:

            served_start = customer.served_start.replace(tzinfo=None)
            served_at = customer.served_at.replace(tzinfo=None)

            ws.cell(row=row, column=1, value=customer.id)
            ws.cell(row=row, column=2, value=customer.ticket_number)
            ws.cell(row=row, column=3, value=(served_at - served_start).total_seconds())
            ws.cell(row=row, column=4, value=served_start.strftime('%H:%M:%S'))  
            ws.cell(row=row, column=5, value=served_at.strftime('%H:%M:%S'))
            ws.cell(row=row, column=6, value=(customer.served_start - customer.created_at).total_seconds())
            ws.cell(row=row, column=7, value=customer.window.number)
            ws.cell(row=row, column=8, value=customer.pasport)
            ws.cell(row=row, column=9, value=customer.category)
            ws.cell(row=row, column=10, value=customer.operator.username)
            ws.cell(row=row, column=11, value=customer.created_at.strftime('%Y-%m-%d'))
            ws.cell(row=row, column=12, value=customer.queue.branch.name)

            for column in range(1, 11):
                column_letter = get_column_letter(column)
                cell = ws[column_letter + str(row)]
                ws.column_dimensions[column_letter].width = max(ws.column_dimensions[column_letter].width, len(str(cell.value)))

            row += 1


        for row in ws.iter_rows(min_row=1, max_row=ws.max_row):
            for cell in row:
                cell.alignment = Alignment(horizontal='center')


        output = BytesIO()
        wb.save(output)
        output.seek(0)


        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename=customers_day_report_{date}.xlsx'
        response.write(output.getvalue())

        return response
    

    @action(detail=False, methods=['get'])
    @swagger_auto_schema(manual_parameters=[
        openapi.Parameter(
            name='start_date',
            in_=openapi.IN_QUERY,
            type=openapi.TYPE_STRING,
            format=openapi.FORMAT_DATE,
            required=True,
            description='Date in the format "YYYY-MM-DD"'
        ),
        openapi.Parameter(
            name='end_date',
            in_=openapi.IN_QUERY,
            type=openapi.TYPE_STRING,
            format=openapi.FORMAT_DATE,
            required=True,
            description='Date in the format "YYYY-MM-DD"'
        )
    ])
    def get_days_report(self, request):

        start_date = request.query_params.get('start_date')
        end_date = request.query_params.get('end_date')


        customers = self.get_customers_by_days(start_date, end_date)


        wb = Workbook()
        ws = wb.active
        ws.title = 'Hourly Report'


        ws['A1'] = 'ID'
        ws['B1'] = 'Номер талона'
        ws['C1'] = 'Время обслуживания(с)'
        ws['D1'] = 'Начало обслуживания'
        ws['E1'] = 'Конец обслуживания'
        ws['F1'] = 'Время ожидания'
        ws['G1'] = 'Номер окна'
        ws['H1'] = 'Номер паспорта'
        ws['I1'] = 'Категория'
        ws['J1'] = 'Оператор'
        ws['K1'] = 'Дата'
        ws['L1'] = 'Филиал'


        row = 2
        for customer in customers:

            served_start = customer.served_start.replace(tzinfo=None)
            served_at = customer.served_at.replace(tzinfo=None)

            ws.cell(row=row, column=1, value=customer.id)
            ws.cell(row=row, column=2, value=customer.ticket_number)
            ws.cell(row=row, column=3, value=(served_at - served_start).total_seconds())
            ws.cell(row=row, column=4, value=served_start.strftime('%H:%M:%S')) 
            ws.cell(row=row, column=5, value=served_at.strftime('%H:%M:%S'))
            ws.cell(row=row, column=6, value=(customer.served_start - customer.created_at).total_seconds())
            ws.cell(row=row, column=7, value=customer.window.number)
            ws.cell(row=row, column=8, value=customer.pasport)
            ws.cell(row=row, column=9, value=customer.category)
            ws.cell(row=row, column=10, value=customer.operator.username)
            ws.cell(row=row, column=11, value=customer.created_at.strftime('%Y-%m-%d'))
            ws.cell(row=row, column=12, value=customer.queue.branch.name)

            for column in range(1, 11):
                column_letter = get_column_letter(column)
                cell = ws[column_letter + str(row)]
                ws.column_dimensions[column_letter].width = max(ws.column_dimensions[column_letter].width, len(str(cell.value)))

            row += 1


        for row in ws.iter_rows(min_row=1, max_row=ws.max_row):
            for cell in row:
                cell.alignment = Alignment(horizontal='center')


        output = BytesIO()
        wb.save(output)
        output.seek(0)

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename=customers_day_report_{start_date}_{end_date}.xlsx'
        response.write(output.getvalue())

        return response
    

    def get_customers_by_days(self, start_date, end_date):
        customers = Customer.objects.filter(created_at__date__range=(start_date, end_date), is_served=True)
        return customers

    def get_customers_by_day(self, date):
        customers = Customer.objects.filter(created_at__date=date, is_served=True)
        return customers
    
    def get_customers_by_year(self, year):
        customers = Customer.objects.filter(created_at__year=year, is_served=True)
        return customers


import pandas as pd
from django.shortcuts import get_object_or_404
from matplotlib.backends.backend_agg import FigureCanvasAgg as FigureCanvas
from rest_framework import status
from rest_framework.response import Response


class OperatorHourReportPDFView(APIView):
    def get_reports_by_hour(self, operator, date, hour):
        reports = OperatorDailyReport.objects.filter(operator=operator, date=date, hour=hour)
        return reports

    def get(self, request, operator_id, date, hour):
        operator = get_object_or_404(User, pk=operator_id)
        reports = self.get_reports_by_hour(operator, date, hour)

        if not reports:
            return Response({'error': 'No reports available for the specified date and hour.'}, status=status.HTTP_404_NOT_FOUND)

        x = ['Max Time', 'Min Time', 'Average Time']
        width = 0.4

        max_time = reports.values_list('max_time', flat=True)
        min_time = reports.values_list('min_time', flat=True)
        average_time = reports.values_list('average_time', flat=True)

        hours = range(len(x))

        plt.figure(figsize=(8, 6))

        plt.bar(hours, max_time, width=width, align='center', label='Max Time')
        plt.bar([h + width for h in hours], min_time, width=width, align='center', label='Min Time')
        plt.bar([h + 2 * width for h in hours], average_time, width=width, align='center', label='Average Time')

        plt.xlabel('Hour')
        plt.ylabel('Time (seconds)')
        plt.title(f'Operator Hourly Time Report - {operator.username} - {date} - Hour {hour}')
        plt.xticks(hours, x)
        plt.legend()

        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename=operator_hour_report_{operator.username}_{date}_hour{hour}.pdf'

        canvas = FigureCanvas(plt.gcf())
        canvas.print_pdf(response)

        return response
    


class OperatorGraphicReportWordViewSet(ViewSet):
    @action(detail=True, methods=['get'])
    @swagger_auto_schema(manual_parameters=[
        openapi.Parameter(
            name='year',
            in_=openapi.IN_QUERY,
            type=openapi.TYPE_STRING,
            format=openapi.FORMAT_DATE,
            required=True,
            description='Date in the format "YYYY"'
        )
    ])
    def get_monthly_average_time(self, request, pk=None):
        operator = User.objects.get(pk=pk)
        year = request.query_params.get('year')

        try:
            year = int(year)
        except ValueError:
            return Response({'error': 'Invalid year format'}, status=status.HTTP_400_BAD_REQUEST)

        start_date = datetime(year, 1, 1)  # Use datetime instead of datetime.date
        end_date = start_date + relativedelta(years=1) - timedelta(days=1)

        customers = Customer.objects.filter(operator=operator, served_start__range=(start_date, end_date))

        average_times = [0] * 12

        for month in range(1, 13):
            month_start = datetime(year, month, 1)  # Use datetime instead of datetime.date
            month_end = month_start + relativedelta(months=1) - timedelta(days=1)
            month_customers = customers.filter(served_start__range=(month_start, month_end))

            if month_customers.exists():
                average_time = sum((customer.served_at - customer.served_start).total_seconds() / 60 for customer in month_customers if customer.served_at and customer.served_start) / len(month_customers)
            else:
                average_time = 0

            average_times[month-1] = average_time

        
        x = ['January', 'February', 'March', 'April', 'May', 'June', 'July', 'August', 'September', 'October', 'November', 'December']
        y = average_times

        plt.figure(figsize=(8, 6))
        plt.plot(x, y)

        plt.xlabel('Month')
        plt.ylabel('Average Time (minutes)')
        plt.title(f'Operator Yearly Average Time Report - {operator.username} - {year}')

        plt.savefig('graph.png')  

        
        doc = Document()
        doc.add_paragraph(f'Operator: {operator.username}')
        doc.add_paragraph(f'Year: {year}')
        doc.add_picture('graph.png', width=Inches(6), height=Inches(4))

        
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename=operator_yearly_report_{operator.username}_{year}.docx'
        doc.save(response)

        return response





class OperatorStatisticReportPDFViewSet(ViewSet):

    @action(detail=True, methods=['get'])
    @swagger_auto_schema(manual_parameters=[
        openapi.Parameter(
            name='year',
            in_=openapi.IN_QUERY,
            type=openapi.TYPE_INTEGER,
            format=openapi.FORMAT_INT64,
            required=True,
            description='Date in the format "YYYY"')
    ])
    def get_yearly_report(self, request, pk=None):
        operator = User.objects.get(pk=pk)
        year = request.query_params.get('year')

        try:
            year = int(year)
        except ValueError:
            return Response({'error': 'Invalid year format'}, status=status.HTTP_400_BAD_REQUEST)
        
        start_date = datetime.date(year, 1, 1)
        end_date = start_date + relativedelta(years=1) - datetime.timedelta(days=1)

        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename=operator_year_report_{operator.username}_{year}.pdf'

        with tempfile.NamedTemporaryFile(suffix='.png') as tmp_file1, \
            tempfile.NamedTemporaryFile(suffix='.png') as tmp_file2:

            plt.figure(figsize=(8, 6))
            x = ['January', 'February', 'March', 'April', 'May', 'June', 'July', 'August', 'September', 'October', 'November', 'December']
            width = 0.25

            total_served = [0] * 12
            max_time = [0] * 12
            min_time = [0] * 12
            average_time = [0] * 12

            for month in range(1, 13):
                month_start = datetime.date(year, month, 1)
                month_end = month_start + relativedelta(months=1) - datetime.timedelta(days=1)
                reports = OperatorDailyReport.objects.filter(operator=operator, date__range=(month_start, month_end))

                if reports.exists():
                    total_served[month-1] = sum(report.total_served for report in reports)
                    max_time[month-1] = max(report.max_time for report in reports)
                    min_time[month-1] = min(report.min_time for report in reports)
                    average_time[month-1] = sum(report.average_time for report in reports) / len(reports)
                else:
                    total_served[month-1] = 0
                    max_time[month-1] = 0
                    min_time[month-1] = 0
                    average_time[month-1] = 0

            months = range(len(x))

            plt.bar([m - width for m in months], max_time, width=width, align='center', label='Max Time')

            plt.bar(months, min_time, width=width, align='edge', label='Min Time')

            plt.bar([m + width for m in months], average_time, width=width, align='center', label='Average Time')

            plt.xlabel('Months')
            plt.ylabel('Time (seconds)')
            plt.title(f'Operator Yearly Time Report - {operator.username}')
            plt.xticks(months, x, rotation=45)
            plt.legend()

            plt.savefig(tmp_file1.name, format='png')
            plt.close()

            plt.figure(figsize=(8, 6))

            plt.bar(months, total_served, width=width, align='center', label='Total Served')

            plt.xlabel('Months')
            plt.ylabel('Number of Servings')
            plt.title(f'Operator Yearly Service Report - {operator.username}')
            plt.xticks(months, x, rotation=45)
            plt.legend()

            plt.savefig(tmp_file2.name, format='png')
            plt.close()

            p = canvas.Canvas(response, pagesize=letter)
            p.setFont('Helvetica', 12)

            p.drawString(100, 100, f'Operator: {operator.username}')
            p.drawString(100, 120, f'Year: {year}')

            p.drawImage(tmp_file1.name, 100, 200, width=400, height=300)
            p.drawImage(tmp_file2.name, 100, 500, width=400, height=300)

            p.showPage()
            p.save()

        return response

    @action(detail=True, methods=['get'])
    @swagger_auto_schema(manual_parameters=[
        openapi.Parameter(
            name='date',
            in_=openapi.IN_QUERY,
            type=openapi.TYPE_STRING,
            format=openapi.FORMAT_DATE,
            required=True,
            description='Date in the format "YYYY-MM-DD"'
        )
    ])
    def get_daily_report(self, request, pk=None):
        operator = User.objects.get(pk=pk)
        date_str = request.query_params.get('date')

        try:
            date = datetime.strptime(date_str, '%Y-%m-%d').date()
        except ValueError:
            return Response({'error': 'Invalid date format'}, status=status.HTTP_400_BAD_REQUEST)
        
        start_date = datetime.combine(date, datetime.min.time())
        end_date = start_date + timedelta(days=1)

        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename=operator_daily_report_{operator.username}_{date}.pdf'

        with tempfile.NamedTemporaryFile(suffix='.png') as tmp_file:

            plt.figure(figsize=(8, 6))
            width = 0.35

            customers = Customer.objects.filter(operator=operator, served_start__range=(start_date, end_date))

            if customers.exists():
                min_time = min((customer.served_at - customer.served_start).total_seconds() / 60 for customer in customers)
                max_time = max((customer.served_at - customer.served_start).total_seconds() / 60 for customer in customers)
                average_time = sum((customer.served_at - customer.served_start).total_seconds() / 60 for customer in customers) / len(customers)
                num_services = customers.count()
            else:
                min_time = 0
                max_time = 0
                average_time = 0
                num_services = 0

            plt.bar(['Min Time', 'Max Time', 'Average Time'], [min_time, max_time, average_time], width=width, align='center')

            plt.xlabel('Time')
            plt.ylabel('Duration (minutes)')
            plt.title(f'Operator Daily Time Report - {operator.username} - {date}')
            plt.ylim(0, max_time * 1.2)

            plt.bar(['Number of Services'], [num_services], width=width, align='center')

            plt.savefig(tmp_file.name, format='png')
            plt.close()

            p = canvas.Canvas(response, pagesize=letter)
            p.setFont('Helvetica', 12)

            p.drawString(100, 100, f'Operator: {operator.username}')
            p.drawString(100, 120, f'Date: {date}')

            p.drawImage(tmp_file.name, 100, 200, width=400, height=300)

            p.showPage()
            p.save()

        return response



class OperatorServedReportExcelViewSet(ViewSet):

    @action(detail=True, methods=['get'])
    @swagger_auto_schema(manual_parameters=[
        openapi.Parameter(
            name='date',
            in_=openapi.IN_QUERY,
            type=openapi.TYPE_STRING,
            format=openapi.FORMAT_DATE,
            required=True,
            description='Date in the format "YYYY-MM-DD"'
        ),
        openapi.Parameter(
            name='start_hour',
            in_=openapi.IN_QUERY,
            type=openapi.TYPE_INTEGER,
            required=True,
            description='Start hour (0-23)'
        ),
        openapi.Parameter(
            name='end_hour',
            in_=openapi.IN_QUERY,
            type=openapi.TYPE_INTEGER,
            required=True,
            description='End hour (0-23)'
        )
    ])
    def get_hour_report(self, request, pk=None):

        operator_id = pk
        date = request.query_params.get('date')
        start_hour = request.query_params.get('start_hour')
        end_hour = request.query_params.get('end_hour')


        operator = get_object_or_404(User, pk=operator_id)
        customers = self.get_customers_by_hour(operator, date, start_hour, end_hour)

        wb = Workbook()
        ws = wb.active
        ws.title = 'Hourly Report'

        ws['A1'] = 'ID'
        ws['B1'] = 'Номер талона'
        ws['C1'] = 'Время обслуживания(с)'
        ws['D1'] = 'Начало обслуживания'
        ws['E1'] = 'Конец обслуживания'
        ws['F1'] = 'Время ожидания'
        ws['G1'] = 'Номер окна'
        ws['H1'] = 'Номер паспорта'
        ws['I1'] = 'Категория'
        ws['J1'] = 'Оператор'
        ws['K1'] = 'Дата'
        ws['L1'] = 'Филиал'

        row = 2
        for customer in customers:

            served_start = customer.served_start.replace(tzinfo=None)
            served_at = customer.served_at.replace(tzinfo=None)

            ws.cell(row=row, column=1, value=customer.id)
            ws.cell(row=row, column=2, value=customer.ticket_number)
            ws.cell(row=row, column=3, value=(served_at - served_start).total_seconds())
            ws.cell(row=row, column=4, value=served_start.strftime('%H:%M:%S'))  # Форматирование времени
            ws.cell(row=row, column=5, value=served_at.strftime('%H:%M:%S'))
            ws.cell(row=row, column=6, value=(customer.served_start - customer.created_at).total_seconds())
            ws.cell(row=row, column=7, value=customer.window.number)
            ws.cell(row=row, column=8, value=customer.pasport)
            ws.cell(row=row, column=9, value=customer.category)
            ws.cell(row=row, column=10, value=customer.operator.username)
            ws.cell(row=row, column=11, value=customer.created_at.strftime('%Y-%m-%d'))
            ws.cell(row=row, column=12, value=customer.queue.branch.name)

            for column in range(1, 11):
                column_letter = get_column_letter(column)
                cell = ws[column_letter + str(row)]
                ws.column_dimensions[column_letter].width = max(ws.column_dimensions[column_letter].width, len(str(cell.value)))

            row += 1


        for row in ws.iter_rows(min_row=1, max_row=ws.max_row):
            for cell in row:
                cell.alignment = Alignment(horizontal='center')


        output = BytesIO()
        wb.save(output)
        output.seek(0)


        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename=operator_hour_report_{operator.username}_{date}_{start_hour}-{end_hour}.xlsx'
        response.write(output.getvalue())

        return response
    

    @action(detail=True, methods=['get'])
    @swagger_auto_schema(manual_parameters=[
        openapi.Parameter(
            name='date',
            in_=openapi.IN_QUERY,
            type=openapi.TYPE_STRING,
            format=openapi.FORMAT_DATE,
            required=True,
            description='Date in the format "YYYY-MM-DD"'
        )
    ])
    def get_day_report(self, request, pk=None):

        operator_id = pk
        date = request.query_params.get('date')


        operator = get_object_or_404(User, pk=operator_id)
        customers = self.get_customers_by_day(operator, date)


        wb = Workbook()
        ws = wb.active
        ws.title = 'Hourly Report'


        ws['A1'] = 'ID'
        ws['B1'] = 'Номер талона'
        ws['C1'] = 'Время обслуживания(с)'
        ws['D1'] = 'Начало обслуживания'
        ws['E1'] = 'Конец обслуживания'
        ws['F1'] = 'Время ожидания'
        ws['G1'] = 'Номер окна'
        ws['H1'] = 'Номер паспорта'
        ws['I1'] = 'Категория'
        ws['J1'] = 'Оператор'
        ws['K1'] = 'Дата'
        ws['L1'] = 'Филиал'


        row = 2
        for customer in customers:

            served_start = customer.served_start.replace(tzinfo=None)
            served_at = customer.served_at.replace(tzinfo=None)

            ws.cell(row=row, column=1, value=customer.id)
            ws.cell(row=row, column=2, value=customer.ticket_number)
            ws.cell(row=row, column=3, value=(served_at - served_start).total_seconds())
            ws.cell(row=row, column=4, value=served_start.strftime('%H:%M:%S'))  
            ws.cell(row=row, column=5, value=served_at.strftime('%H:%M:%S'))
            ws.cell(row=row, column=6, value=(customer.served_start - customer.created_at).total_seconds())
            ws.cell(row=row, column=7, value=customer.window.number)
            ws.cell(row=row, column=8, value=customer.pasport)
            ws.cell(row=row, column=9, value=customer.category)
            ws.cell(row=row, column=10, value=customer.operator.username)
            ws.cell(row=row, column=11, value=customer.created_at.strftime('%Y-%m-%d'))
            ws.cell(row=row, column=12, value=customer.queue.branch.name)

            for column in range(1, 11):
                column_letter = get_column_letter(column)
                cell = ws[column_letter + str(row)]
                ws.column_dimensions[column_letter].width = max(ws.column_dimensions[column_letter].width, len(str(cell.value)))

            row += 1


        for row in ws.iter_rows(min_row=1, max_row=ws.max_row):
            for cell in row:
                cell.alignment = Alignment(horizontal='center')


        output = BytesIO()
        wb.save(output)
        output.seek(0)


        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename=operator_day_report_{operator.username}_{date}.xlsx'
        response.write(output.getvalue())

        return response
    

    @action(detail=True, methods=['get'])
    @swagger_auto_schema(manual_parameters=[
        openapi.Parameter(
            name='start_date',
            in_=openapi.IN_QUERY,
            type=openapi.TYPE_STRING,
            format=openapi.FORMAT_DATE,
            required=True,
            description='Date in the format "YYYY-MM-DD"'
        ),
        openapi.Parameter(
            name='end_date',
            in_=openapi.IN_QUERY,
            type=openapi.TYPE_STRING,
            format=openapi.FORMAT_DATE,
            required=True,
            description='Date in the format "YYYY-MM-DD"'
        )
    ])
    def get_days_report(self, request, pk=None):

        operator_id = pk
        start_date = request.query_params.get('start_date')
        end_date = request.query_params.get('end_date')


        operator = get_object_or_404(User, pk=operator_id)
        customers = self.get_customers_by_days(operator, start_date, end_date)


        wb = Workbook()
        ws = wb.active
        ws.title = 'Hourly Report'


        ws['A1'] = 'ID'
        ws['B1'] = 'Номер талона'
        ws['C1'] = 'Время обслуживания(с)'
        ws['D1'] = 'Начало обслуживания'
        ws['E1'] = 'Конец обслуживания'
        ws['F1'] = 'Время ожидания'
        ws['G1'] = 'Номер окна'
        ws['H1'] = 'Номер паспорта'
        ws['I1'] = 'Категория'
        ws['J1'] = 'Оператор'
        ws['K1'] = 'Дата'
        ws['L1'] = 'Филиал'


        row = 2
        for customer in customers:

            served_start = customer.served_start.replace(tzinfo=None)
            served_at = customer.served_at.replace(tzinfo=None)

            ws.cell(row=row, column=1, value=customer.id)
            ws.cell(row=row, column=2, value=customer.ticket_number)
            ws.cell(row=row, column=3, value=(served_at - served_start).total_seconds())
            ws.cell(row=row, column=4, value=served_start.strftime('%H:%M:%S')) 
            ws.cell(row=row, column=5, value=served_at.strftime('%H:%M:%S'))
            ws.cell(row=row, column=6, value=(customer.served_start - customer.created_at).total_seconds())
            ws.cell(row=row, column=7, value=customer.window.number)
            ws.cell(row=row, column=8, value=customer.pasport)
            ws.cell(row=row, column=9, value=customer.category)
            ws.cell(row=row, column=10, value=customer.operator.username)
            ws.cell(row=row, column=11, value=customer.created_at.strftime('%Y-%m-%d'))
            ws.cell(row=row, column=12, value=customer.queue.branch.name)

            for column in range(1, 11):
                column_letter = get_column_letter(column)
                cell = ws[column_letter + str(row)]
                ws.column_dimensions[column_letter].width = max(ws.column_dimensions[column_letter].width, len(str(cell.value)))

            row += 1


        for row in ws.iter_rows(min_row=1, max_row=ws.max_row):
            for cell in row:
                cell.alignment = Alignment(horizontal='center')


        output = BytesIO()
        wb.save(output)
        output.seek(0)

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename=operator_day_report_{operator.username}_{start_date}_{end_date}.xlsx'
        response.write(output.getvalue())

        return response
    

    @action(detail=True, methods=['get'])
    @swagger_auto_schema(manual_parameters=[
        openapi.Parameter(
            name='year',
            in_=openapi.IN_QUERY,
            type=openapi.TYPE_STRING,
            format=openapi.FORMAT_INT64,
            required=True,
            description='Date in the format "YYYY"')
    ])
    def get_year_report(self, request, pk=None):

        operator_id = pk
        year = request.query_params.get('year')


        operator = get_object_or_404(User, pk=operator_id)
        customers = self.get_customers_by_year(operator, year)


        wb = Workbook()
        ws = wb.active
        ws.title = 'Hourly Report'


        ws['A1'] = 'ID'
        ws['B1'] = 'Номер талона'
        ws['C1'] = 'Время обслуживания(с)'
        ws['D1'] = 'Начало обслуживания'
        ws['E1'] = 'Конец обслуживания'
        ws['F1'] = 'Время ожидания'
        ws['G1'] = 'Номер окна'
        ws['H1'] = 'Номер паспорта'
        ws['I1'] = 'Категория'
        ws['J1'] = 'Оператор'
        ws['K1'] = 'Дата'
        ws['L1'] = 'Филиал'

   
        row = 2
        for customer in customers: 

            served_start = customer.served_start.replace(tzinfo=None)
            served_at = customer.served_at.replace(tzinfo=None)

            ws.cell(row=row, column=1, value=customer.id)
            ws.cell(row=row, column=2, value=customer.ticket_number)
            ws.cell(row=row, column=3, value=(served_at - served_start).total_seconds())
            ws.cell(row=row, column=4, value=served_start.strftime('%H:%M:%S')) 
            ws.cell(row=row, column=5, value=served_at.strftime('%H:%M:%S'))
            ws.cell(row=row, column=6, value=(customer.served_start - customer.created_at).total_seconds())
            ws.cell(row=row, column=7, value=customer.window.number)
            ws.cell(row=row, column=8, value=customer.pasport)
            ws.cell(row=row, column=9, value=customer.category)
            ws.cell(row=row, column=10, value=customer.operator.username)
            ws.cell(row=row, column=11, value=customer.created_at.strftime('%Y-%m-%d'))
            ws.cell(row=row, column=12, value=customer.queue.branch.name)

            for column in range(1, 11):
                column_letter = get_column_letter(column)
                cell = ws[column_letter + str(row)]
                ws.column_dimensions[column_letter].width = max(ws.column_dimensions[column_letter].width, len(str(cell.value)))

            row += 1

      
        for row in ws.iter_rows(min_row=1, max_row=ws.max_row):
            for cell in row:
                cell.alignment = Alignment(horizontal='center')

        
        output = BytesIO()
        wb.save(output)
        output.seek(0)

       
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename=operator_day_report_{operator.username}_{year}.xlsx'
        response.write(output.getvalue())

        return response


    def get_customers_by_hour(self, operator, date, start_hour, end_hour):
        customers = Customer.objects.filter(operator=operator, served_start__date=date, served_start__hour__range=(start_hour, end_hour), is_served=True)
        return customers
    
    def get_customers_by_days(self, operator, start_date, end_date):
        customers = Customer.objects.filter(operator=operator, created_at__date__range=(start_date, end_date), is_served=True)
        return customers

    def get_customers_by_day(self, operator, date):
        customers = Customer.objects.filter(operator=operator, created_at__date=date, is_served=True)
        return customers
    
    def get_customers_by_year(self, operator, year):
        customers = Customer.objects.filter(operator=operator, created_at__year=year, is_served=True)
        return customers
    


class OperatorCancelledReportExcelViewSet(ViewSet):
    @action(detail=True, methods=['get'])
    @swagger_auto_schema(manual_parameters=[
        openapi.Parameter(
            name='date',
            in_=openapi.IN_QUERY,
            type=openapi.TYPE_STRING,
            format=openapi.FORMAT_DATE,
            required=True,
            description='Date in the format "YYYY-MM-DD"'
        )
    ])
    def get_day_report(self, request, pk=None):

        operator_id = pk
        date = request.query_params.get('date')


        operator = get_object_or_404(User, pk=operator_id)
        customers = self.get_customers_by_day(operator, date)

        
        wb = Workbook()
        ws = wb.active
        ws.title = 'Hourly Report'

        
        ws['A1'] = 'ID'
        ws['B1'] = 'Номер талона'
        ws['C1'] = 'Номер паспорта'
        ws['D1'] = 'Категория'
        ws['E1'] = 'Дата'
        ws['F1'] = 'Филиал'


        
        row = 2
        for customer in customers:

            ws.cell(row=row, column=1, value=customer.id)
            ws.cell(row=row, column=2, value=customer.ticket_number)
            ws.cell(row=row, column=3, value=customer.pasport)
            ws.cell(row=row, column=4, value=customer.category)
            ws.cell(row=row, column=5, value=customer.created_at.strftime('%Y-%m-%d'))
            ws.cell(row=row, column=6, value=customer.queue.branch.name)

            for column in range(1, 5):
                column_letter = get_column_letter(column)
                cell = ws[column_letter + str(row)]
                ws.column_dimensions[column_letter].width = max(ws.column_dimensions[column_letter].width, len(str(cell.value)))

            row += 1

       
        for row in ws.iter_rows(min_row=1, max_row=ws.max_row):
            for cell in row:
                cell.alignment = Alignment(horizontal='center')

        
        output = BytesIO()
        wb.save(output)
        output.seek(0)

      
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename=operator_day_report_{operator.username}_{date}.xlsx'
        response.write(output.getvalue())

        return response
    

    @action(detail=True, methods=['get'])
    @swagger_auto_schema(manual_parameters=[
        openapi.Parameter(
            name='start_date',
            in_=openapi.IN_QUERY,
            type=openapi.TYPE_STRING,
            format=openapi.FORMAT_DATE,
            required=True,
            description='Date in the format "YYYY-MM-DD"'
        ),
        openapi.Parameter(
            name='end_date',
            in_=openapi.IN_QUERY,
            type=openapi.TYPE_STRING,
            format=openapi.FORMAT_DATE,
            required=True,
            description='Date in the format "YYYY-MM-DD"'
        )
    ])
    def get_days_report(self, request, pk=None):

        operator_id = pk
        start_date = request.query_params.get('start_date')
        end_date = request.query_params.get('end_date')


        operator = get_object_or_404(User, pk=operator_id)
        customers = self.get_customers_by_days(operator, start_date, end_date)

        
        wb = Workbook()
        ws = wb.active
        ws.title = 'Hourly Report'

        
        ws['A1'] = 'ID'
        ws['B1'] = 'Номер талона'
        ws['C1'] = 'Номер паспорта'
        ws['D1'] = 'Категория'
        ws['E1'] = 'Дата'
        ws['F1'] = 'Филиал'


     
        row = 2
        for customer in customers:

            ws.cell(row=row, column=1, value=customer.id)
            ws.cell(row=row, column=2, value=customer.ticket_number)
            ws.cell(row=row, column=3, value=customer.pasport)
            ws.cell(row=row, column=4, value=customer.category)
            ws.cell(row=row, column=5, value=customer.created_at.strftime('%Y-%m-%d'))
            ws.cell(row=row, column=6, value=customer.queue.branch.name)

            for column in range(1, 5):
                column_letter = get_column_letter(column)
                cell = ws[column_letter + str(row)]
                ws.column_dimensions[column_letter].width = max(ws.column_dimensions[column_letter].width, len(str(cell.value)))

            row += 1

        
        for row in ws.iter_rows(min_row=1, max_row=ws.max_row):
            for cell in row:
                cell.alignment = Alignment(horizontal='center')

        
        output = BytesIO()
        wb.save(output)
        output.seek(0)

       
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename=operator_day_report_{operator.username}_{start_date}_{end_date}.xlsx'
        response.write(output.getvalue())

        return response
    

    @action(detail=True, methods=['get'])
    @swagger_auto_schema(manual_parameters=[
        openapi.Parameter(
            name='year',
            in_=openapi.IN_QUERY,
            type=openapi.TYPE_STRING,
            format=openapi.FORMAT_INT64,
            required=True,
            description='Date in the format "YYYY"')
    ])
    def get_year_report(self, request, pk=None):

        operator_id = pk
        year = request.query_params.get('year')


        operator = get_object_or_404(User, pk=operator_id)
        customers = self.get_customers_by_year(operator, year)

      
        wb = Workbook()
        ws = wb.active
        ws.title = 'Hourly Report'

       
        ws['A1'] = 'ID'
        ws['B1'] = 'Номер талона'
        ws['C1'] = 'Номер паспорта'
        ws['D1'] = 'Категория'
        ws['E1'] = 'Дата'
        ws['F1'] = 'Филиал'


      
        row = 2
        for customer in customers:

            ws.cell(row=row, column=1, value=customer.id)
            ws.cell(row=row, column=2, value=customer.ticket_number)
            ws.cell(row=row, column=3, value=customer.pasport)
            ws.cell(row=row, column=4, value=customer.category)
            ws.cell(row=row, column=5, value=customer.created_at.strftime('%Y-%m-%d'))
            ws.cell(row=row, column=6, value=customer.queue.branch.name)

            for column in range(1, 5):
                column_letter = get_column_letter(column)
                cell = ws[column_letter + str(row)]
                ws.column_dimensions[column_letter].width = max(ws.column_dimensions[column_letter].width, len(str(cell.value)))

            row += 1

      
        for row in ws.iter_rows(min_row=1, max_row=ws.max_row):
            for cell in row:
                cell.alignment = Alignment(horizontal='center')

        
        output = BytesIO()
        wb.save(output)
        output.seek(0)

        
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename=operator_day_report_{operator.username}_{year}.xlsx'
        response.write(output.getvalue())

        return response
    

    def get_customers_by_days(self, operator, start_date, end_date):
        customers = Customer.objects.filter(operator=operator, created_at__date__range=(start_date, end_date), is_served=False)
        return customers

    def get_customers_by_day(self, operator, date):
        customers = Customer.objects.filter(operator=operator, created_at__date=date, is_served=False)
        return customers
    
    def get_customers_by_year(self, operator, year):
        customers = Customer.objects.filter(operator=operator, created_at__year=year, is_served=False)
        return customers